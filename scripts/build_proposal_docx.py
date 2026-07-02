#!/usr/bin/env python3
"""Generate the Flowstack partnership proposal as a .docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for lvl, size in [("Heading 1", 16), ("Heading 2", 13)]:
    st = doc.styles[lvl]
    st.font.color.rgb = ACCENT
    st.font.size = Pt(size)
    st.font.name = "Calibri"


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def para(text="", italic=False, color=None, size=None, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def table(headers, rows, bold_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            if bold_col is not None and i == bold_col:
                r.bold = True
    doc.add_paragraph()
    return t


# ---- Logo header ----
LOGO = "/home/theone/automation_dashboard/docs/assets/flowstack-logo.png"
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.add_run().add_picture(LOGO, width=Inches(2.4))

word_p = doc.add_paragraph()
word_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
wr = word_p.add_run("FLOWSTACK")
wr.bold = True
wr.font.size = Pt(20)
wr.font.color.rgb = ACCENT
# letter-spacing via the XML (python-docx has no direct API)
from docx.oxml.ns import qn
rPr = wr._element.get_or_add_rPr()
spc = rPr.makeelement(qn("w:spacing"), {qn("w:val"): "60"})
rPr.append(spc)

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tg = tag_p.add_run("An AI agent for your team. Live in 60 seconds.")
tg.italic = True
tg.font.size = Pt(10)
tg.font.color.rgb = GREY

# thin divider
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
dv = div.add_run("─" * 30)
dv.font.color.rgb = ACCENT
doc.add_paragraph()

# ---- Title block ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Partnership Proposal")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = ACCENT

for label, value in [
    ("Prepared by", "Neoptolemos Papadiofantous, Founder"),
    ("For", "[Partner Name]"),
    ("Date", "27 June 2026"),
    ("Status", "Draft for discussion"),
]:
    p = doc.add_paragraph()
    rb = p.add_run(f"{label}:  ")
    rb.bold = True
    rb.font.color.rgb = GREY
    p.add_run(value)

doc.add_paragraph()

# ---- 1 ----
heading("1. Executive Summary")
para("Flowstack is a built, live, revenue-ready product. Over the past ~6–10 months I have "
     "designed, engineered, and funded the entire platform and infrastructure single-handedly. "
     "The one thing the business needs now is sales reach — and that's where you come in.")
para("This proposal sets out a partnership that is fair to both sides by construction: I keep "
     "ownership of what I have already built and continue to run it; you are rewarded — generously "
     "and immediately — for the revenue you actually bring in. We start with a low-risk trial, and "
     "your stake in the business grows as the revenue you generate grows.")

# ---- 2 ----
heading("2. Roles & Contributions")
table(
    ["", "Founder (me)", "Partner (you)"],
    [
        ["Owns", "Product, engineering, infrastructure, technical ops, roadmap",
         "Sales, outbound, lead generation, closing, account management"],
        ["Contributed to date", "~6–10 months of full-time build + all infra/AI costs, funded personally",
         "Joining now; no capital or build contribution yet"],
        ["Contributes going forward", "Maintains and develops the product; carries running costs",
         "Drives all new revenue"],
    ],
    bold_col=0,
)

# ---- 3 ----
heading("3. The Product You'd Be Selling")
para("A high-margin SaaS with clear, proven pricing:")
table(
    ["Plan", "Price (EUR)", "Monthly credits", "Gross margin"],
    [
        ["Starter", "€99 / mo", "2,500", "88–96%"],
        ["Operator", "€399 / mo", "25,000", "71–89%"],
    ],
    bold_col=0,
)
para("Plus one-off top-up packs (€29 / €119 / €399) and custom top-ups (€10–€2,000).")
p = doc.add_paragraph()
p.add_run("Why this matters for you: ").bold = True
p.add_run("direct AI/infrastructure cost is only ~4–29% of revenue. The business runs at high "
          "margin — which is exactly what makes a strong sales commission affordable and sustainable.")

# ---- 4 ----
heading("4. How You Get Paid")
heading("Phase 1 — Commission (Months 1–3, trial)", level=2)
p = doc.add_paragraph()
p.add_run("You earn ").italic = False
p.add_run("25% of net revenue").bold = True
p.add_run(" (revenue after direct AI/infra/payment costs) on every customer you originate, for as "
          "long as that customer stays subscribed. No equity, no lock-in — pure performance.")
para("What that looks like per customer (at a conservative ~75% net margin):")
table(
    ["You close…", "Your recurring commission"],
    [
        ["1 Starter (€99/mo)", "~€21 / month"],
        ["1 Operator (€399/mo)", "~€75 / month"],
        ["1 Large top-up (€399, one-off)", "~€75 one-off"],
    ],
    bold_col=1,
)
para("At scale:")
table(
    ["Scenario", "MRR", "Your cut / mo", "Founder keeps / mo"],
    [
        ["10 Starter + 3 Operator", "€2,187", "€410", "€1,230"],
        ["20 Starter + 10 Operator", "€5,970", "€1,119", "€3,359"],
        ["40 Starter + 25 Operator", "€13,935", "€2,613", "€7,838"],
    ],
    bold_col=2,
)
para("Commission compounds: every retained customer keeps paying you month after month.")

heading("Phase 2 — Earned Equity (after a successful trial)", level=2)
para("Once you've proven you can sell, we convert your role into a real ownership stake:")
bullet("Up to 15–25% equity, vesting against cumulative revenue milestones — earned by results, not time served.")
bullet("1-year cliff: nothing vests if you stop producing or leave within the first 12 months.")
bullet("Founder retains a controlling majority (≥ 65–70%), reflecting the existing build, IP, and infrastructure.")

# ---- 5 ----
heading("5. Revenue & Profit Handling")
numbered("All gross revenue flows into the business account.")
numbered("Direct costs are deducted first — AI/provider costs, infrastructure, Stripe fees, taxes.")
numbered("Your commission is paid on net revenue, protecting us both from unprofitable deals "
         "(e.g. a customer running the premium model tier on the cheapest plan).")
numbered("Remaining profit is reinvested or distributed per the ownership split.")

# ---- 6 ----
heading("6. Protecting What Already Exists")
bullet("All existing IP, code, infrastructure, and the product itself remain 100% the Founder's as "
       "the founding contribution. This is non-negotiable and reflects the months and costs already invested.")
bullet("Any equity earned is in the go-forward business — not a retroactive claim on what has already been built.")
bullet("Cost recovery: before any profit is split, the business reimburses the Founder for documented "
       "infra/AI costs already incurred, from early revenue.")

# ---- 7 ----
heading("7. Trial & Exit Terms")
bullet("3-month commission-only trial. No equity, no commitment. If you produce, we formalise and move "
       "to Phase 2. If not, we part cleanly with no entanglement.")
bullet("On departure: vested equity is retained per the schedule; unvested equity is forfeited. "
       "Commissions on existing accounts taper for a defined period (e.g. 6 months) after leaving, then stop.")

# ---- 8 ----
heading("8. Why This Is Fair to Both of Us")
p = doc.add_paragraph(style="List Bullet")
p.add_run("For you: ").bold = True
p.add_run("strong, immediate, recurring income with zero upfront risk, and a clear, earnable path to "
          "real ownership. A confident salesperson should welcome a deal that pays this well on results.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("For me: ").bold = True
p.add_run("I keep what I built and funded, the business never pays out more than the margin can support, "
          "and equity is only ever given for revenue that actually materialises.")

# ---- 9 ----
heading("9. Next Steps")
numbered("Review and discuss these terms together.")
numbered("Agree the trial commission rate and equity ceiling.")
numbered("Have a lawyer paper the final agreement before anything is signed.")
numbered("Set a start date and a 3-month revenue target to review against.")

# ---- footer note ----
doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run(
    "Figures are based on Flowstack's internal pricing audit (11 June 2026) and assume a conservative "
    "~75% blended net margin; actual net per customer varies with the model tier they run. This document "
    "is a commercial proposal, not legal advice — any binding agreement should be reviewed by a qualified lawyer."
)
nr.italic = True
nr.font.size = Pt(9)
nr.font.color.rgb = GREY

# ---- signature block ----
doc.add_paragraph()
sig = doc.add_table(rows=2, cols=2)
sig.autofit = True
labels = [("Founder", "Neoptolemos Papadiofantous"), ("Partner", "[Partner Name]")]
for col, (role, name) in enumerate(labels):
    c = sig.cell(0, col)
    c.paragraphs[0].add_run(f"\n\n_______________________________\n").bold = False
    r = c.add_paragraph().add_run(f"{role}: {name}")
    r.bold = True
    sig.cell(1, col).paragraphs[0].add_run("Date: ____________________")

# ---- Running header + footer ----
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement

section = doc.sections[0]

# header: small logo, right aligned
hdr_p = section.header.paragraphs[0]
hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hdr_p.add_run().add_picture(LOGO, width=Inches(0.9))

# footer: confidentiality left, page number right (tab-separated)
ftr_p = section.footer.paragraphs[0]
fr = ftr_p.add_run("Confidential — Flowstack Partnership Proposal")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY
ftr_p.add_run("\t\t")
# PAGE field
def add_page_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(_qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(_qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(_qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(8); run.font.color.rgb = GREY

pr = ftr_p.add_run("Page ")
pr.font.size = Pt(8); pr.font.color.rgb = GREY
add_page_field(ftr_p)

out = "/home/theone/automation_dashboard/docs/Flowstack-Partnership-Proposal.docx"
doc.save(out)
print("Saved:", out)
